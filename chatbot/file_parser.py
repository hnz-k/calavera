import PyPDF2
import docx
from typing import Optional


class FileParser:
    """Parser untuk ekstraksi teks dari berbagai format file"""
    
    @staticmethod
    def parse_txt(file_path: str) -> str:
        """Parse TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # Fallback ke latin-1 jika UTF-8 gagal
            with open(file_path, 'r', encoding='latin-1') as f:
                return f.read()
    
    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Parse PDF file"""
        try:
            text = ""
            with open(file_path, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                
                # Ekstrak teks dari semua halaman
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            if not text.strip():
                return "⚠️ File PDF ini tidak mengandung teks yang bisa dibaca (mungkin gambar scan)."
            
            return text.strip()
            
        except Exception as e:
            return f"⚠️ Gagal membaca PDF: {str(e)}"
    
    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Parse DOCX file"""
        try:
            doc = docx.Document(file_path)
            text = ""
            
            # Ekstrak teks dari semua paragraf
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Ekstrak teks dari tabel (jika ada)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                return "⚠️ File DOCX ini kosong atau tidak mengandung teks."
            
            return text.strip()
            
        except Exception as e:
            return f"⚠️ Gagal membaca DOCX: {str(e)}"
    
    @staticmethod
    def parse_file(file_path: str, file_extension: str) -> Optional[str]:
        """
        Main parser function - deteksi format dan parse sesuai tipe
        
        Args:
            file_path: Path lengkap ke file
            file_extension: Extension file (pdf, docx, txt, dll)
        
        Returns:
            Teks hasil ekstraksi atau error message
        """
        extension = file_extension.lower()
        
        if extension == 'txt':
            return FileParser.parse_txt(file_path)
        
        elif extension == 'pdf':
            return FileParser.parse_pdf(file_path)
        
        elif extension in ['docx', 'doc']:
            return FileParser.parse_docx(file_path)
        
        else:
            return f"⚠️ Format file .{extension} belum didukung."